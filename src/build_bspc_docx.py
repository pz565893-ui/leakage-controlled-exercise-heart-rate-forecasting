"""Build an editable, single-column BSPC manuscript draft from project Markdown."""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from docx_table_geometry import apply_bspc_table_borders, apply_table_geometry


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "main_manuscript.md"
BIB = ROOT / "references" / "references.bib"
OUTPUT = ROOT / "manuscript" / "BSPC_main_manuscript_draft.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
FONT = "Calibri"


def set_run_font(run, *, size: float | None = None, color=None, bold=None, italic=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(9.5)
    caption.font.bold = False
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.line_spacing = 1.15

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("BSPC manuscript draft | 23 July 2026")
    set_run_font(hr, size=8.5, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = fp.add_run("Page ")
    set_run_font(label, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    set_run_font(run, size=9, color=GRAY)
    run._r.extend([fld_char1, instr, fld_char2])

    props = doc.core_properties
    props.title = (
        "Uncertainty-Aware Exercise Heart-Rate Forecasting under User and Sport "
        "Distribution Shifts: A Leakage-Controlled Multi-Dataset Study"
    )
    props.subject = "Editable manuscript draft for Biomedical Signal Processing and Control"
    props.author = ""
    props.keywords = "heart-rate forecasting; wearable sensors; distribution shift"


def add_title_block(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, size=19, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Full paper | Biomedical Signal Processing and Control")
    set_run_font(r, size=10.5, color=GRAY, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("[Author names and affiliations to be supplied]")
    set_run_font(r, size=11, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Corresponding author: [name, postal address, and email to be supplied]")
    set_run_font(r, size=10, color=GRAY)


def clean_tex(text: str) -> str:
    text = text.replace("\\ldots", "...")
    text = text.replace("\\(", "").replace("\\)", "")
    text = text.replace("\\[", "").replace("\\]", "")
    text = text.replace("\\text{", "").replace("\\mathrm{", "")
    text = text.replace("\\in", " in ").replace("\\le", " <= ")
    text = text.replace("\\alpha", "alpha")
    text = text.replace("--", "–")
    text = text.replace("\\", "")
    return text


def add_inline_markup(paragraph, text: str, *, default_size: float | None = None):
    text = clean_tex(text)
    token = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+`)")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=default_size)
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif value.startswith("*"):
            run = paragraph.add_run(value[1:-1])
            set_run_font(run, size=default_size, italic=True)
        else:
            run = paragraph.add_run(value[1:-1])
            set_run_font(run, size=default_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size)


def parse_bibtex(text: str) -> dict[str, dict[str, str]]:
    entries: dict[str, dict[str, str]] = {}
    chunks = re.split(r"(?m)(?=^@)", text)
    for chunk in chunks:
        header = re.match(r"@(\w+)\{([^,]+),", chunk)
        if not header:
            continue
        key = header.group(2).strip()
        fields: dict[str, str] = {"entry_type": header.group(1).lower()}
        for line in chunk.splitlines()[1:]:
            match = re.match(r"\s*(\w+)\s*=\s*\{(.*)\},?\s*$", line)
            if match:
                fields[match.group(1).lower()] = match.group(2).strip()
        entries[key] = fields
    return entries


_LATEX_ACCENT_PATTERN = re.compile(
    r"\{\\(?P<accent>['`^\"~=\.uvHckr])(?:\{(?P<braced>[A-Za-z])\}|(?P<plain>[A-Za-z]))\}"
)
_LATEX_COMBINING = {
    "'": "\u0301",
    "`": "\u0300",
    "^": "\u0302",
    '"': "\u0308",
    "~": "\u0303",
    "=": "\u0304",
    ".": "\u0307",
    "u": "\u0306",
    "v": "\u030c",
    "H": "\u030b",
    "c": "\u0327",
    "k": "\u0328",
    "r": "\u030a",
}


def decode_latex_accents(value: str) -> str:
    """Convert braced BibTeX accent commands to composed Unicode text."""

    def replace(match: re.Match) -> str:
        letter = match.group("braced") or match.group("plain")
        return unicodedata.normalize("NFC", letter + _LATEX_COMBINING[match.group("accent")])

    return _LATEX_ACCENT_PATTERN.sub(replace, value)


def clean_bib_value(value: str) -> str:
    value = decode_latex_accents(value)
    return value.replace("{", "").replace("}", "").replace("--", "–").strip()


def format_author(name: str) -> str:
    name = clean_bib_value(name)
    if "," not in name:
        return name
    family, given = [part.strip() for part in name.split(",", 1)]
    initials = "".join(part[0] for part in re.findall(r"[^\W\d_]+", given) if part)
    return f"{family} {initials}"


def format_reference(fields: dict[str, str]) -> str:
    authors = fields.get("author", "")
    author_text = ", ".join(format_author(item) for item in authors.split(" and "))
    title = clean_bib_value(fields.get("title", ""))
    venue = clean_bib_value(
        fields.get("journal") or fields.get("booktitle") or fields.get("publisher", "")
    )
    year = clean_bib_value(fields.get("year", ""))
    volume = clean_bib_value(fields.get("volume", ""))
    number = clean_bib_value(fields.get("number", ""))
    pages = clean_bib_value(fields.get("pages", ""))
    doi = clean_bib_value(fields.get("doi", ""))
    url = clean_bib_value(fields.get("url", ""))

    dataset_marker = "[dataset] " if fields.get("entry_type") == "dataset" else ""
    ref = f"{dataset_marker}{author_text}. {title}. {venue}. {year}"
    if volume:
        ref += f";{volume}"
        if number:
            ref += f"({number})"
    if pages:
        ref += f":{pages}"
    ref += "."
    if doi:
        ref += f" https://doi.org/{doi}."
    elif url:
        ref += f" {url}."
    return re.sub(r"\s+", " ", ref).replace("..", ".")


def replace_citations(text: str) -> tuple[str, list[str]]:
    order: list[str] = []

    def repl(match: re.Match) -> str:
        content = match.group(1)
        keys = re.findall(r"@([A-Za-z0-9_:-]+)", content)
        if not keys:
            return match.group(0)
        nums = []
        for key in keys:
            if key not in order:
                order.append(key)
            nums.append(str(order.index(key) + 1))
        return "[" + ",".join(nums) + "]"

    converted = re.sub(r"\[([^\]]*@[^\]]+)\]", repl, text)
    return converted, order


def add_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_markdown_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Normal Table"
    for r_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(values):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline_markup(p, value, default_size=9)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
        if r_idx == 0:
            add_repeat_table_header(table.rows[0])
    widths = [6240] + [1040] * (len(rows[0]) - 1)
    if sum(widths) != 9360:
        widths[-1] += 9360 - sum(widths)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    apply_bspc_table_borders(table)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_reference_list(doc: Document, citation_order: list[str], entries):
    for idx, key in enumerate(citation_order, start=1):
        if key not in entries:
            raise KeyError(f"Citation key missing from BibTeX: {key}")
        p = doc.add_paragraph()
        # Reference entries are intentionally ragged-right. Justification can
        # produce conspicuous inter-word gaps in DOI- and URL-heavy lines.
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.10
        run = p.add_run(f"[{idx}] ")
        set_run_font(run, size=9, bold=True)
        run = p.add_run(format_reference(entries[key]))
        set_run_font(run, size=9)


def remove_separate_caption_section(lines: list[str]) -> list[str]:
    """Exclude captions from the main Word source when a caption DOCX is supplied."""
    output: list[str] = []
    skipping = False
    for line in lines:
        if line.rstrip() == "## Figure captions":
            skipping = True
            continue
        if skipping and line.startswith("## "):
            skipping = False
        if not skipping:
            output.append(line)
    return output


def build():
    raw = SOURCE.read_text(encoding="utf-8")
    converted, citation_order = replace_citations(raw)
    bib_entries = parse_bibtex(BIB.read_text(encoding="utf-8"))

    title = converted.splitlines()[0].removeprefix("# ").strip()
    lines = remove_separate_caption_section(converted.splitlines())
    start = lines.index("## Abstract")

    doc = Document()
    configure_document(doc)
    add_title_block(doc, title)

    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if line == "## References":
            doc.add_heading("References", level=1)
            add_reference_list(doc, citation_order, bib_entries)
            break
        if not line.strip():
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
            rows = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
                i += 1
            headers = [cell.strip() for cell in line.strip().strip("|").split("|")]
            add_markdown_table(doc, [headers] + rows)
            continue
        if line.startswith("### "):
            doc.add_heading(clean_tex(line[4:].strip()), level=2)
        elif line.startswith("## "):
            doc.add_heading(clean_tex(line[3:].strip()), level=1)
        elif line.startswith("# "):
            doc.add_heading(clean_tex(line[2:].strip()), level=1)
        elif re.match(r"^\d+\.\s+", line):
            # Academic override: protocol items are real subparagraphs rather than fake lists.
            content = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            add_inline_markup(p, content)
        elif line.startswith("**Table ") or line.startswith("**Fig. ") or line.startswith("**Supplementary Fig. "):
            p = doc.add_paragraph(style="Caption")
            if line.startswith("**Table "):
                p.paragraph_format.keep_with_next = True
            add_inline_markup(p, line, default_size=9.5)
        elif line.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            add_inline_markup(p, line)
        else:
            p = doc.add_paragraph()
            add_inline_markup(p, line)
        i += 1

    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"citation_order={len(citation_order)}; bib_entries={len(bib_entries)}")


if __name__ == "__main__":
    build()
