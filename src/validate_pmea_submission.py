"""Validate the Physiological Measurement-facing manuscript package."""

from __future__ import annotations

import argparse
import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_MAIN = ROOT / "manuscript" / "PMEA_complete_document_draft.docx"
DEFAULT_SUPPLEMENT = ROOT / "manuscript" / "PMEA_supplementary_material.docx"
DEFAULT_SOURCE = ROOT / "manuscript" / "main_manuscript.md"
DEFAULT_OUTPUT = ROOT / "outputs" / "audit" / "PMEA_SUBMISSION_VALIDATION.json"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS = {"w": W_NS, "wp": WP_NS}
EXPECTED_TITLE = (
    "Boundary-dependent reliability of exercise heart-rate forecasts across users, "
    "sports, and data sources: a leakage-controlled study"
)


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'’-]+\b", text, flags=re.UNICODE))


def markdown_section(text: str, start: str, end: str) -> str:
    match = re.search(
        rf"(?ms)^{re.escape(start)}\s*$\n(.*?)(?=^{re.escape(end)}(?:\s|$))",
        text,
    )
    if not match:
        raise ValueError(f"Missing section boundary: {start} -> {end}")
    return match.group(1).strip()


def docx_alt_text(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    return [
        (node.get("descr") or "").strip()
        for node in root.findall(".//wp:docPr", NS)
        if (node.get("descr") or "").strip()
    ]


def _paragraph_index(paragraphs: list[Any], text: str) -> int:
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == text:
            return index
    raise ValueError(f"Paragraph not found: {text}")


def validate_main(path: Path, source_path: Path) -> dict[str, Any]:
    source = source_path.read_text(encoding="utf-8")
    abstract = markdown_section(source, "## Abstract", "**Keywords:**")
    main_body = markdown_section(source, "## 1. Introduction", "## References")
    document = Document(path)
    paragraphs = document.paragraphs
    texts = [paragraph.text.strip() for paragraph in paragraphs if paragraph.text.strip()]
    joined = "\n".join(texts)

    abstract_words = word_count(abstract)
    main_words = word_count(main_body)
    structured_labels = ["Objective.", "Approach.", "Main results.", "Significance."]
    structured_present = {
        label: any(text.startswith(label) for text in texts) for label in structured_labels
    }

    reference_start = _paragraph_index(paragraphs, "References") + 1
    references = [
        paragraph.text.strip()
        for paragraph in paragraphs[reference_start:]
        if paragraph.text.strip()
    ]
    reference_initials = [
        re.sub(r"[^A-Za-zÀ-ÖØ-öø-ÿ].*$", "", reference).lower()
        for reference in references
    ]
    captions = [
        paragraph.text.strip()
        for paragraph in paragraphs
        if paragraph.style.name == "Caption" and paragraph.text.strip().startswith("Figure ")
    ]
    alt_text = docx_alt_text(path)
    citation_markers = re.findall(r"\[@|(?<![\w.])@[A-Za-z0-9_:-]+", joined)
    numeric_citations = re.findall(r"\[(?:\d+[–,;\s]*)+\]", joined)
    placeholders = sorted(
        set(
            re.findall(
                r"\[[^\]\n]*(?:to be supplied|to be confirmed)[^\]\n]*\]",
                joined,
                flags=re.IGNORECASE,
            )
        )
    )
    narrative = source.partition("## Figure captions")[0]
    first_figure_citations = [narrative.find(f"Fig. {number}") for number in range(1, 5)]

    checks = {
        "target_title_present": EXPECTED_TITLE in texts,
        "author_names_present": all(
            author in joined for author in ("PANG KEREN", "MIN CHANGRONG")
        ),
        "author_affiliations_present": all(
            affiliation in joined
            for affiliation in (
                "Department of Sports & Health Science, Shinhan University",
                "Criminal Investigation Police University of China",
            )
        ),
        "corresponding_author_contact_present": all(
            value in joined
            for value in (
                "20248657@o.shinhan.ac.kr",
                "0009-0007-2506-9206",
                "mcr19940816@gmail.com",
            )
        ),
        "funding_statement_finalized": (
            "This research did not receive any specific grant" in joined
        ),
        "competing_interest_statement_finalized": (
            "The authors declare no competing interests." in joined
        ),
        "legacy_target_journal_absent": "Biomedical Signal Processing and Control" not in joined,
        "abstract_at_most_250_words": abstract_words <= 250,
        "structured_abstract_complete": all(structured_present.values()),
        "main_text_at_most_8000_words": main_words <= 8000,
        "harvard_citations_no_raw_keys": not citation_markers,
        "harvard_citations_no_numeric_brackets": not numeric_citations,
        "references_alphabetized": reference_initials == sorted(reference_initials),
        "references_include_titles": len(references) >= 30,
        "four_embedded_figures": len(document.inline_shapes) == 4,
        "four_figure_captions": len(captions) == 4,
        "all_figures_have_alt_text": len(alt_text) == 4,
        "figure_first_citations_are_ordered": all(
            position >= 0 for position in first_figure_citations
        )
        and first_figure_citations == sorted(first_figure_citations),
        "figure_captions_are_ordered": [
            re.match(r"Figure (\d+)\.", caption).group(1)
            for caption in captions
            if re.match(r"Figure (\d+)\.", caption)
        ]
        == ["1", "2", "3", "4"],
        "ai_disclosure_appears_once": source.count("OpenAI Codex desktop") == 1,
        "internal_ai_reminder_absent": "Before submission, all authors must inspect" not in source,
        "percentage_range_typography_correct": source.count("2.7%–13.9%") == 2
        and "2.7%--13.9%" not in source,
        "main_table_present": len(document.tables) >= 1,
        "structured_section_headings_present": all(
            heading in texts
            for heading in (
                "Data availability",
                "Code availability",
                "Ethics statement",
                "CRediT authorship contribution statement",
                "Funding",
                "Declaration of competing interest",
                "Acknowledgements",
                "References",
            )
        ),
    }
    return {
        "path": path.resolve().relative_to(ROOT.resolve()).as_posix(),
        "checks": checks,
        "pass": all(checks.values()),
        "abstract_words": abstract_words,
        "main_text_words": main_words,
        "structured_abstract_labels": structured_present,
        "reference_count": len(references),
        "figure_count": len(document.inline_shapes),
        "figure_caption_count": len(captions),
        "figure_alt_text_count": len(alt_text),
        "table_count": len(document.tables),
        "author_release_placeholders": placeholders,
    }


def validate_supplement(path: Path) -> dict[str, Any]:
    document = Document(path)
    texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    joined = "\n".join(texts)
    low_support_rows: list[list[str]] = []
    for table in document.tables:
        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            if values and values[0].lower().startswith("joint"):
                if any(value in {"18", "19", "20"} for value in values):
                    low_support_rows.append(values)

    alt_text = docx_alt_text(path)
    major_table_numbers = [
        int(match.group(1))
        for text in texts
        if (match := re.fullmatch(r"Table S(\d+)\..*", text))
    ]
    checks = {
        "target_title_present": EXPECTED_TITLE in texts,
        "author_names_present": all(
            author in joined for author in ("PANG KEREN", "MIN CHANGRONG")
        ),
        "target_journal_correct": (
            "Target journal: Physiological Measurement" in joined
            and "Biomedical Signal Processing and Control" not in joined
        ),
        "recorded_gender_outcome_section_omitted": not re.search(
            r"recorded-(?:female|gender)|female minus male", joined, re.IGNORECASE
        ),
        "supplementary_major_tables_are_contiguous": major_table_numbers
        == list(range(1, 19)),
        "recorded_gender_figure_panel_omitted": not re.search(
            r"recorded-(?:female|gender)|female minus male", joined, re.IGNORECASE
        ),
        "joint_outcomes_obey_30_user_threshold": not low_support_rows,
        "one_embedded_supplementary_figure": len(document.inline_shapes) == 1,
        "supplementary_figure_has_alt_text": len(alt_text) == 1,
        "threshold_disclosure_present": "at least 30 users" in joined,
    }
    return {
        "path": path.resolve().relative_to(ROOT.resolve()).as_posix(),
        "checks": checks,
        "pass": all(checks.values()),
        "figure_count": len(document.inline_shapes),
        "figure_alt_text_count": len(alt_text),
        "table_count": len(document.tables),
        "major_table_numbers": major_table_numbers,
        "low_support_joint_rows": low_support_rows,
    }


def validate_package(
    main_path: Path,
    supplement_path: Path,
    source_path: Path,
    output_path: Path,
) -> dict[str, Any]:
    main = validate_main(main_path, source_path)
    supplement = validate_supplement(supplement_path)
    author_blockers = main["author_release_placeholders"]
    audit = {
        "validator_version": "1.0.0",
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "target_journal": "Physiological Measurement",
        "scientific_and_format_pass": main["pass"] and supplement["pass"],
        "submission_ready": main["pass"] and supplement["pass"] and not author_blockers,
        "status": (
            "READY"
            if main["pass"] and supplement["pass"] and not author_blockers
            else "READY_EXCEPT_AUTHOR_AND_RELEASE_INPUTS"
            if main["pass"] and supplement["pass"]
            else "FAIL"
        ),
        "main": main,
        "supplement": supplement,
        "author_and_release_blockers": author_blockers,
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8")
    return audit


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--main", type=Path, default=DEFAULT_MAIN)
    parser.add_argument("--supplement", type=Path, default=DEFAULT_SUPPLEMENT)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    audit = validate_package(args.main, args.supplement, args.source, args.output)
    print(json.dumps(audit, indent=2, ensure_ascii=False))
    return 0 if audit["scientific_and_format_pass"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
