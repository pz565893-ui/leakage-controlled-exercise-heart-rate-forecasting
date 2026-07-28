"""Build a complete Physiological Measurement review manuscript.

The initial-submission document uses Harvard author-year citations, an
alphabetized reference list containing article titles, and figures/tables
embedded near the corresponding text. Author and institutional placeholders
remain explicit until supplied by the authors.
"""

from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_bspc_docx import (
    add_inline_markup,
    add_markdown_table,
    clean_bib_value,
    clean_tex,
    format_author,
    parse_bibtex,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "main_manuscript.md"
BIB = ROOT / "references" / "references.bib"
OUTPUT = ROOT / "manuscript" / "PMEA_complete_document_draft.docx"

FIGURES = {
    1: ROOT / "figures" / "Figure_1_study_design.png",
    2: ROOT / "figures" / "Figure_2_primary_performance.png",
    3: ROOT / "figures" / "Figure_3_sport_shift_PMEA.png",
    4: ROOT / "figures" / "Figure_4_uncertainty_calibration.png",
}

FIGURE_ALT = {
    1: (
        "Study-design schematic showing a five-minute past-only context, a "
        "current-workout and completed-history model, and separated temporal, "
        "unseen-user, held-sport, joint-shift, and frozen cross-source tests."
    ),
    2: (
        "Four-panel comparison of hierarchical heart-rate forecast MAE across "
        "strict temporal, unseen-user, and frozen cross-source evaluations, "
        "with paired user-bootstrap effects for history versus zero-history training."
    ),
    3: (
        "Four-panel held-sport analysis showing point error, model-minus-EWMA "
        "differences, and user support. Joint-shift outcome cells with fewer "
        "than 30 users are marked not reported."
    ),
    4: (
        "Six-panel assessment of empirical interval coverage, interval width, "
        "weighted interval score, and width-error association across internal "
        "and frozen cross-source evaluations."
    ),
}

INSERT_BEFORE = {
    "## 2. Materials and methods": 1,
    "### 3.3. Frozen cross-source evaluation showed lower natural-mix accuracy": 2,
    "### 3.5. Empirical interval coverage was near nominal internally but lower across sources": 3,
    "### 3.6. Ablations and sensitivity analyses bounded the main claims": 4,
}

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)


def configure_document(doc: Document) -> None:
    """Apply a restrained narrative-manuscript style and explicit geometry."""

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9.5)
    caption.font.bold = False
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Physiological Measurement research paper draft | 28 July 2026")
    set_run_font(run, size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer.add_run("Page ")
    set_run_font(label, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = footer.add_run()
    set_run_font(field_run, size=9, color=GRAY)
    field_run._r.extend([begin, instruction, end])

    props = doc.core_properties
    props.title = (
        "Boundary-dependent reliability of exercise heart-rate forecasts across "
        "users, sports, and data sources: a leakage-controlled study"
    )
    props.subject = "Research paper draft for Physiological Measurement"
    props.author = "PANG KEREN; MIN CHANGRONG"
    props.keywords = "heart-rate forecasting; physiological measurement; distribution shift"


def add_title_block(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=19, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Research paper | Physiological Measurement")
    set_run_font(run, size=10.5, color=GRAY, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("PANG KEREN")
    set_run_font(run, size=11, bold=True)
    run = p.add_run("1,*")
    set_run_font(run, size=8, bold=True)
    run.font.superscript = True
    run = p.add_run(" and MIN CHANGRONG")
    set_run_font(run, size=11, bold=True)
    run = p.add_run("2")
    set_run_font(run, size=8, bold=True)
    run.font.superscript = True

    affiliations = (
        (
            "1",
            "Department of Sports & Health Science, Shinhan University, 95 Hoam-ro, "
            "Uijeongbu-si, Gyeonggi-do 11644, Republic of Korea",
        ),
        (
            "2",
            "Criminal Investigation Police University of China, 83 Tawan Street, "
            "Huanggu District, Shenyang, Liaoning 110854, China",
        ),
    )
    for number, affiliation in affiliations:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(number)
        set_run_font(run, size=7.5)
        run.font.superscript = True
        run = p.add_run(f" {affiliation}")
        set_run_font(run, size=9.5, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(
        "*Corresponding author: PANG KEREN; 20248657@o.shinhan.ac.kr; "
        "ORCID 0009-0007-2506-9206"
    )
    set_run_font(run, size=9.5, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Co-author email: MIN CHANGRONG; mcr19940816@gmail.com")
    set_run_font(run, size=9.5, color=GRAY)


def family_name(author: str) -> str:
    value = clean_bib_value(author).strip()
    if not value:
        return "Unknown"
    if "," in value:
        return value.split(",", 1)[0].strip()
    return value.split()[-1]


def author_list(fields: dict[str, str]) -> list[str]:
    return [item.strip() for item in fields.get("author", "").split(" and ") if item.strip()]


def citation_author(fields: dict[str, str]) -> str:
    authors = author_list(fields)
    if not authors:
        return clean_bib_value(fields.get("publisher") or fields.get("title", "Unknown"))
    surnames = [family_name(item) for item in authors]
    if len(surnames) == 1:
        return surnames[0]
    if len(surnames) == 2:
        return f"{surnames[0]} and {surnames[1]}"
    return f"{surnames[0]} et al"


def make_citation_labels(
    keys: list[str], entries: dict[str, dict[str, str]]
) -> dict[str, str]:
    base = {
        key: f"{citation_author(entries[key])} {clean_bib_value(entries[key].get('year', 'n.d.'))}"
        for key in keys
    }
    counts = Counter(base.values())
    labels: dict[str, str] = {}
    for label, count in counts.items():
        group = [key for key in keys if base[key] == label]
        group.sort(key=lambda key: clean_bib_value(entries[key].get("title", "")).lower())
        for index, key in enumerate(group):
            suffix = chr(ord("a") + index) if count > 1 else ""
            labels[key] = f"{label}{suffix}"
    return labels


def replace_citations(
    text: str, entries: dict[str, dict[str, str]]
) -> tuple[str, list[str], dict[str, str]]:
    cited = []
    for key in re.findall(r"(?<![\w.])@([A-Za-z0-9_:-]+)", text):
        if key not in entries:
            raise KeyError(f"Citation key missing from BibTeX: {key}")
        if key not in cited:
            cited.append(key)
    labels = make_citation_labels(cited, entries)

    def repl(match: re.Match[str]) -> str:
        keys = re.findall(r"(?<![\w.])@([A-Za-z0-9_:-]+)", match.group(1))
        if not keys:
            return match.group(0)
        return "(" + "; ".join(labels[key] for key in keys) + ")"

    converted = re.sub(r"\[([^\]]*@[A-Za-z0-9_:-]+[^\]]*)\]", repl, text)
    return converted, cited, labels


def normalize_submission_typography(text: str) -> str:
    """Keep negative numbers together and avoid a stranded ASCII minus in Word."""

    # The Markdown source uses a double hyphen for numerical ranges. Resolve
    # those ranges before converting standalone negative signs.
    text = re.sub(r"(?<=[A-Za-z0-9%])--(?=[A-Za-z0-9(])", "–", text)
    return re.sub(r"(?<![\w])-(?=\d)", "−", text)


def format_authors(fields: dict[str, str]) -> str:
    formatted = [format_author(item) for item in author_list(fields)]
    if not formatted:
        return clean_bib_value(fields.get("publisher") or "Unknown")
    if len(formatted) == 1:
        return formatted[0]
    if len(formatted) == 2:
        return f"{formatted[0]} and {formatted[1]}"
    return ", ".join(formatted[:-1]) + f" and {formatted[-1]}"


def format_reference(fields: dict[str, str], label: str) -> str:
    authors = format_authors(fields)
    year = re.search(r"(\d{4}|n\.d\.)[a-z]?$", label)
    year_label = year.group(0) if year else clean_bib_value(fields.get("year", "n.d."))
    title = clean_bib_value(fields.get("title", ""))
    venue = clean_bib_value(
        fields.get("journal") or fields.get("booktitle") or fields.get("publisher", "")
    )
    volume = clean_bib_value(fields.get("volume", ""))
    number = clean_bib_value(fields.get("number", ""))
    pages = clean_bib_value(fields.get("pages", ""))
    doi = clean_bib_value(fields.get("doi", ""))
    url = clean_bib_value(fields.get("url", ""))

    ref = f"{authors} {year_label} {title}"
    if venue:
        ref += f" {venue}"
    if volume:
        ref += f" {volume}"
        if number:
            ref += f"({number})"
    if pages:
        ref += f" {pages}"
    if doi:
        ref += f" https://doi.org/{doi}"
    elif url:
        ref += f" {url}"
    return re.sub(r"\s+", " ", ref).strip()


def reference_sort_key(
    key: str, entries: dict[str, dict[str, str]], labels: dict[str, str]
) -> tuple[str, str, str]:
    fields = entries[key]
    authors = author_list(fields)
    first = family_name(authors[0]) if authors else citation_author(fields)
    return (
        first.lower(),
        labels[key].lower(),
        clean_bib_value(fields.get("title", "")).lower(),
    )


def add_reference_list(
    doc: Document,
    cited: list[str],
    entries: dict[str, dict[str, str]],
    labels: dict[str, str],
) -> None:
    for key in sorted(cited, key=lambda item: reference_sort_key(item, entries, labels)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.10
        run = p.add_run(format_reference(entries[key], labels[key]))
        set_run_font(run, size=9.5)


def extract_captions(lines: list[str]) -> tuple[list[str], dict[int, str]]:
    content: list[str] = []
    captions: dict[int, str] = {}
    in_captions = False
    for line in lines:
        if line.rstrip() == "## Figure captions":
            in_captions = True
            continue
        if in_captions and line.startswith("## "):
            in_captions = False
        if in_captions:
            caption_line = line.strip()
            match = re.match(
                r"^\*\*Fig\.\s+(\d+)\.\s*(.+?)\*\*\s*(.*)$", caption_line
            )
            if match:
                captions[int(match.group(1))] = (
                    "**Figure "
                    + match.group(1)
                    + ". "
                    + match.group(2)
                    + "** "
                    + match.group(3)
                )
            continue
        content.append(line)
    return content, captions


def add_figure(doc: Document, number: int, caption: str) -> None:
    path = FIGURES[number]
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.25))
    shape._inline.docPr.set("descr", FIGURE_ALT[number])
    shape._inline.docPr.set("title", f"Figure {number}")

    cp = doc.add_paragraph(style="Caption")
    add_inline_markup(cp, caption, default_size=9.5)


def build() -> None:
    raw = SOURCE.read_text(encoding="utf-8")
    entries = parse_bibtex(BIB.read_text(encoding="utf-8"))
    converted, cited, labels = replace_citations(raw, entries)
    converted = normalize_submission_typography(converted)
    lines, captions = extract_captions(converted.splitlines())
    title = lines[0].removeprefix("# ").strip()
    start = lines.index("## Abstract")

    missing_captions = sorted(set(FIGURES) - set(captions))
    if missing_captions:
        raise ValueError(f"Missing source captions for figures: {missing_captions}")

    doc = Document()
    configure_document(doc)
    add_title_block(doc, title)

    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if line in INSERT_BEFORE:
            number = INSERT_BEFORE[line]
            add_figure(doc, number, captions[number])
        if line == "## References":
            doc.add_heading("References", level=1)
            add_reference_list(doc, cited, entries, labels)
            break
        if not line.strip():
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:|-]+\|$", lines[i + 1]
        ):
            headers = [cell.strip() for cell in line.strip().strip("|").split("|")]
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(
                    [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                )
                i += 1
            add_markdown_table(doc, [headers] + rows)
            continue
        if line.startswith("### "):
            doc.add_heading(clean_tex(line[4:].strip()), level=2)
        elif line.startswith("## "):
            doc.add_heading(clean_tex(line[3:].strip()), level=1)
        elif line.startswith("# "):
            doc.add_heading(clean_tex(line[2:].strip()), level=1)
        elif re.match(r"^\d+\.\s+", line):
            content = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            add_inline_markup(p, content)
        elif line.startswith("**Table "):
            p = doc.add_paragraph(style="Caption")
            p.paragraph_format.keep_with_next = True
            add_inline_markup(p, line, default_size=9.5)
        elif line.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            add_inline_markup(p, line)
        else:
            p = doc.add_paragraph()
            add_inline_markup(p, line)
        i += 1

    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"cited_keys={len(cited)}; bibliography_entries={len(entries)}; figures={len(FIGURES)}")


if __name__ == "__main__":
    build()
