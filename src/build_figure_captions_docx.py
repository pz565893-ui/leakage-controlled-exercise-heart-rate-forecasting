"""Build the separate editable figure-caption file requested by BSPC."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "main_manuscript.md"
OUTPUT = ROOT / "manuscript" / "BSPC_Figure_Captions.docx"
FONT = "Calibri"


def extract_captions(markdown: str) -> list[tuple[str, str]]:
    marker = "## Figure captions"
    if marker not in markdown:
        raise ValueError("figure-caption section is missing")
    section = markdown.split(marker, 1)[1]
    section = section.split("\n## ", 1)[0]
    pattern = re.compile(r"^\*\*(.+?)\*\*\s*(.+)$", re.MULTILINE)
    captions = [(match.group(1).strip(), match.group(2).strip()) for match in pattern.finditer(section)]
    if not captions:
        raise ValueError("no figure captions were found")
    labels = [label for label, _ in captions]
    if len(labels) != len(set(labels)):
        raise ValueError("duplicate figure-caption label")
    return captions


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def build(source: Path = SOURCE, output: Path = OUTPUT) -> Path:
    captions = extract_captions(source.read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("Figure captions")
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(
        "Uncertainty-Aware Exercise Heart-Rate Forecasting under User and Sport "
        "Distribution Shifts: A Leakage-Controlled Multi-Dataset Study"
    )
    run.italic = True
    run.font.name = FONT
    run.font.size = Pt(10)

    for label, text in captions:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_after = Pt(10)
        label_run = paragraph.add_run(label + " ")
        label_run.bold = True
        label_run.font.name = FONT
        body_run = paragraph.add_run(text)
        body_run.font.name = FONT

    properties = document.core_properties
    properties.title = "Figure captions for BSPC submission"
    properties.subject = "Separate editable figure-caption file"
    properties.author = ""
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def main() -> int:
    output = build()
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
