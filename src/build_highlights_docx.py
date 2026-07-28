"""Build the separate Elsevier Highlights Word file from the canonical text."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "highlights.txt"
OUTPUT = ROOT / "manuscript" / "BSPC_Highlights.docx"


def set_font(run, size=11, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def main():
    bullets = [line.strip() for line in SOURCE.read_text(encoding="utf-8").splitlines() if line.strip()]
    if not 3 <= len(bullets) <= 5:
        raise ValueError(f"Elsevier highlights require 3-5 bullets; found {len(bullets)}")
    too_long = [(line, len(line)) for line in bullets if len(line) > 85]
    if too_long:
        raise ValueError(f"Highlights exceed 85 characters: {too_long}")

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    set_font(title.add_run("Highlights"), size=15, bold=True)

    for line in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.15
        set_font(paragraph.add_run(line), size=11)

    props = doc.core_properties
    props.title = "Highlights: leakage-controlled exercise heart-rate forecasting"
    props.subject = "Highlights for Biomedical Signal Processing and Control"
    props.author = ""
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

