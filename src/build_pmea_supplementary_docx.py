"""Build the Physiological Measurement-facing supplementary Word file.

The numerical source remains ``supplementary_material.md``.  This presentation
layer enforces the journal-facing 30-user reporting rule: outcome rows from
joint user-sport cells below that threshold and the recorded-gender outcome
section are omitted.  Tables after that omission are renumbered so the
submission-facing sequence remains continuous.  The underlying audit artifacts
are not changed.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_supplementary_docx as base
from build_bspc_docx import add_inline_markup, set_run_font


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "supplementary_material.md"
OUTPUT = ROOT / "manuscript" / "PMEA_supplementary_material.docx"
FIGURE = ROOT / "figures" / "Supplementary_Figure_1_ablation_sensitivity_PMEA.png"
REPORTING_THRESHOLD = 30


def _cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _integer(value: str) -> int | None:
    match = re.search(r"\d[\d,]*", value)
    return int(match.group(0).replace(",", "")) if match else None


def _is_low_support_row(section: str, line: str) -> bool:
    if not line.startswith("|") or re.match(r"^\|[\s:|-]+\|$", line):
        return False
    cells = _cells(line)
    if not cells or cells[0].lower() == "regime":
        return False

    user_index: int | None = None
    if section.startswith("### Table S5a") and cells[0].lower().startswith("joint"):
        user_index = 6
    elif section.startswith("### Table S5b") and cells[0].lower().startswith("joint"):
        user_index = 9
    elif section.startswith("### Table S5c") and cells[0].lower().startswith("joint"):
        user_index = 4
    elif section.startswith("## Table S17"):
        user_index = 5

    if user_index is None or user_index >= len(cells):
        return False
    users = _integer(cells[user_index])
    return users is not None and users < REPORTING_THRESHOLD


def _renumber_after_gender_omission(text: str) -> str:
    """Map source Table S9--S19 labels to PMEA Table S8--S18."""

    def replace(match: re.Match[str]) -> str:
        return f"S{int(match.group(1)) - 1}{match.group(2)}"

    return re.sub(r"\bS(9|1[0-9])([a-z]?)\b", replace, text)


def filter_for_pmea(lines: list[str]) -> list[str]:
    """Return the submission-facing supplementary presentation."""

    filtered: list[str] = []
    section = ""
    skip_gender = False
    for line in lines:
        if line.startswith("## Table S8."):
            skip_gender = True
            section = line
            continue
        if skip_gender:
            if line.startswith("## Table S9."):
                skip_gender = False
            else:
                continue

        if line.startswith("## ") or line.startswith("### "):
            section = line

        if _is_low_support_row(section, line):
            continue

        replacements = {
            (
                "## Uncertainty-Aware Exercise Heart-Rate Forecasting under User and Sport "
                "Distribution Shifts: A Leakage-Controlled Multi-Dataset Study"
            ): (
                "## Boundary-dependent reliability of exercise heart-rate forecasts across "
                "users, sports, and data sources: a leakage-controlled study"
            ),
            "**Target journal:** *Biomedical Signal Processing and Control*": (
                "**Target journal:** *Physiological Measurement*"
            ),
            "Joint intersections below 25 users are explicitly cautionary.": (
                "Journal-facing joint user–sport outcomes are shown only for cells with at least "
                "30 users; lower-support intersections are retained in the immutable audit trail "
                "but omitted from outcome rows."
            ),
            "Low-support joint intersections remain cautionary.": (
                "Joint user–sport outcome rows are shown only for cells with at least 30 users."
            ),
            (
                "Joint user--sport rows are exploratory regardless of user count, and rows below "
                "25 users receive an additional caution flag."
            ): (
                "Joint user–sport rows remain exploratory; journal-facing outcome rows require at "
                "least 30 users."
            ),
            "Rows with fewer than 25 users remain cautionary.": (
                "Outcome rows with fewer than 30 users are omitted from this journal-facing "
                "supplement; their support counts remain visible in main Figure 3."
            ),
            (
                "**Supplementary Fig. 1. Ablation, stride sensitivity, and subgroup boundaries.** "
                "(a) Reference-seed paired multimodal-minus-HR-only MAE differences. (b) Reference-seed "
                "change in MAE when the frozen unseen-user model is evaluated every 60 s rather than "
                "every 300 s. (c) Five-seed history-informed-minus-zero-history-trained effects, with "
                "paired per-user differences averaged over seeds before user bootstrap. (d) Reference-seed "
                "recorded-female-minus-recorded-male descriptive MAE differences; the unseen-user "
                "recorded-female subgroup contains only 10 users."
            ): (
                "**Supplementary Figure 1. Ablation and reporting-stride sensitivity.** "
                "(a) Reference-seed paired multimodal-minus-HR-only MAE differences. "
                "(b) Reference-seed change in MAE when the frozen unseen-user model is evaluated "
                "every 60 s rather than every 300 s. (c) Five-seed history-informed-minus-zero-history-"
                "trained effects, with paired per-user differences averaged over seeds before user bootstrap."
            ),
            (
                "- Stride sensitivity and recorded-gender contrasts: version 0.15.0 and 0.16.0 result artifacts."
            ): "- Stride sensitivity: version 0.15.0 result artifacts.",
        }
        rewritten = line
        for old, new in replacements.items():
            rewritten = rewritten.replace(old, new)
        filtered.append(_renumber_after_gender_omission(rewritten))
    return filtered


def configure(doc: Document) -> None:
    base.configure(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(
        header.add_run("Physiological Measurement supplementary material | 28 July 2026"),
        size=8,
        color=base.GRAY,
    )
    doc.core_properties.title = (
        "Supplementary material: boundary-dependent reliability of exercise "
        "heart-rate forecasts"
    )
    doc.core_properties.subject = (
        "Physiological Measurement supplementary methods, tables, provenance, and figure"
    )


def add_figure(doc: Document) -> None:
    if not FIGURE.is_file():
        raise FileNotFoundError(FIGURE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(str(FIGURE), width=Inches(6.0))
    shape._inline.docPr.set(
        "descr",
        "Three-panel supplementary figure showing signal ablation, reporting-stride "
        "sensitivity, and completed-history effects.",
    )
    shape._inline.docPr.set("title", "Supplementary Figure 1")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True


def build() -> None:
    lines = filter_for_pmea(SOURCE.read_text(encoding="utf-8").splitlines())
    doc = Document()
    configure(doc)

    i = 0
    figure_added = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:|-]+\|$", lines[i + 1]
        ):
            headers = _cells(line)
            rows = [headers]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(_cells(lines[i]))
                i += 1
            base.add_table(doc, rows)
            continue
        if line == "## Supplementary figure caption" and not figure_added:
            heading = doc.add_heading("Supplementary figure", level=1)
            heading.paragraph_format.page_break_before = True
            heading.paragraph_format.keep_with_next = True
            add_figure(doc)
            figure_added = True
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
            if heading_text.startswith("Boundary-dependent reliability"):
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_markup(
                    paragraph,
                    "**PANG KEREN**¹,* and **MIN CHANGRONG**²",
                )
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            add_inline_markup(paragraph, line[2:])
        elif line.startswith("**Supplementary Fig"):
            paragraph = doc.add_paragraph(style="Caption")
            add_inline_markup(paragraph, line, default_size=8.5)
        else:
            paragraph = doc.add_paragraph()
            add_inline_markup(paragraph, line)
        i += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
