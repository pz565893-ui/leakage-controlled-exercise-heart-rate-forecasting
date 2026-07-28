"""Build a landscape, editable supplementary-material DOCX."""

from __future__ import annotations

import math
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_bspc_docx import add_inline_markup, set_run_font
from docx_table_geometry import apply_bspc_table_borders, apply_table_geometry


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "supplementary_material.md"
OUTPUT = ROOT / "manuscript" / "BSPC_supplementary_material.docx"
FIGURE = ROOT / "figures" / "Supplementary_Figure_1_ablation_sensitivity.png"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
FONT = "Calibri"
CONTENT_DXA = 14400


def configure(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    specs = {
        "Heading 1": (16, BLUE, 12, 7),
        "Heading 2": (13, BLUE, 9, 5),
        "Heading 3": (11, DARK_BLUE, 7, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(8.5)
    caption.font.bold = False
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("BSPC supplementary material | 23 July 2026")
    set_run_font(r, size=8, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Page ")
    set_run_font(r, size=8, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    rr = footer.add_run()
    set_run_font(rr, size=8, color=GRAY)
    rr._r.extend([begin, instr, end])

    props = doc.core_properties
    props.title = "Supplementary material: leakage-controlled exercise heart-rate forecasting"
    props.subject = "Supplementary methods, tables, provenance, and figure"
    props.author = ""


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def column_widths(rows: list[list[str]]) -> list[int]:
    count = len(rows[0])
    max_lengths = []
    for idx in range(count):
        length = max(len(re.sub(r"[*`]", "", row[idx])) for row in rows)
        max_lengths.append(min(max(length, 6), 42))
    weights = [math.sqrt(value) for value in max_lengths]
    minimum = 720 if count >= 10 else 900
    widths = [minimum] * count
    remaining = CONTENT_DXA - minimum * count
    if remaining < 0:
        minimum = CONTENT_DXA // count
        widths = [minimum] * count
        remaining = CONTENT_DXA - sum(widths)
    weight_sum = sum(weights)
    for idx, weight in enumerate(weights):
        widths[idx] += int(round(remaining * weight / weight_sum))
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]):
    # Chain the explanatory paragraph into the table and keep the repeated
    # header with the first data row. This prevents a caption/note/header block
    # from being stranded at the foot of a page with all data on the next page.
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.keep_with_next = True
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Normal Table"
    font_size = 6.5 if len(rows[0]) >= 10 else 7.2 if len(rows[0]) >= 7 else 8.0
    for r_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(values):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline_markup(p, value, default_size=font_size)
            if r_idx == 0:
                p.paragraph_format.keep_with_next = True
                for run in p.runs:
                    run.bold = True
        if r_idx == 0:
            repeat_header(table.rows[0])
    widths = column_widths(rows)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_DXA,
        indent_dxa=80,
        cell_margins_dxa={"top": 60, "bottom": 60, "start": 80, "end": 80},
    )
    apply_bspc_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Keep the heading, figure, and caption on the same landscape page in Word.
    # At 9.5 in the image is taller than the usable page area and Word strands
    # the heading on an otherwise blank page.
    shape = run.add_picture(str(FIGURE), width=Inches(7.4))
    shape._inline.docPr.set(
        "descr",
        "Four-panel supplementary figure showing signal ablation, reporting-stride sensitivity, history effects, and recorded-gender descriptive contrasts.",
    )
    p.paragraph_format.space_after = Pt(4)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure(doc)

    i = 0
    figure_added = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
            headers = [cell.strip() for cell in line.strip().strip("|").split("|")]
            rows = [headers]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, rows)
            continue
        if line == "## Supplementary figure caption" and not figure_added:
            heading = doc.add_heading("Supplementary figure", level=1)
            # Put the break on the heading itself. A standalone break paragraph
            # can be pushed to the next page when the preceding table fills the
            # page, producing an otherwise empty page before the figure.
            heading.paragraph_format.page_break_before = True
            heading.paragraph_format.keep_with_next = True
            add_figure(doc)
            figure_added = True
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            add_inline_markup(p, line[2:])
        elif line.startswith("**Supplementary Fig."):
            p = doc.add_paragraph(style="Caption")
            add_inline_markup(p, line, default_size=8.5)
        else:
            p = doc.add_paragraph()
            add_inline_markup(p, line)
        i += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
