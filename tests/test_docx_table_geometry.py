from __future__ import annotations

import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from docx_table_geometry import (  # noqa: E402
    apply_bspc_table_borders,
    apply_table_geometry,
)


class DocxTableGeometryTests(unittest.TestCase):
    def test_width_is_synchronized_across_table_grid_and_cells(self) -> None:
        table = Document().add_table(rows=2, cols=2)
        apply_table_geometry(table, [3000, 6000], table_width_dxa=9000)

        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        self.assertIsNotNone(table_width)
        self.assertEqual(table_width.get(qn("w:type")), "dxa")
        self.assertEqual(table_width.get(qn("w:w")), "9000")

        grid = [
            int(column.get(qn("w:w")))
            for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        self.assertEqual(grid, [3000, 6000])

        for row in table.rows:
            cell_widths = [
                int(cell._tc.get_or_add_tcPr().find(qn("w:tcW")).get(qn("w:w")))
                for cell in row.cells
            ]
            self.assertEqual(cell_widths, grid)

    def test_declared_table_width_must_match_columns(self) -> None:
        table = Document().add_table(rows=1, cols=2)
        with self.assertRaisesRegex(ValueError, "must sum to table_width_dxa"):
            apply_table_geometry(table, [3000, 6000], table_width_dxa=8000)

    def test_bspc_borders_have_no_vertical_or_inside_rules(self) -> None:
        table = Document().add_table(rows=2, cols=2)
        apply_bspc_table_borders(table)
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        self.assertIsNotNone(borders)
        self.assertEqual(borders.find(qn("w:top")).get(qn("w:val")), "single")
        self.assertEqual(borders.find(qn("w:bottom")).get(qn("w:val")), "single")
        for edge in ("left", "right", "insideH", "insideV"):
            self.assertEqual(borders.find(qn(f"w:{edge}")).get(qn("w:val")), "nil")
        for cell in table.rows[0].cells:
            cell_borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
            self.assertEqual(
                cell_borders.find(qn("w:bottom")).get(qn("w:val")), "single"
            )


if __name__ == "__main__":
    unittest.main()
